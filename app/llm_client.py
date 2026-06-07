import json
import os
import re
from pathlib import Path
from typing import Iterable
from urllib import request

from docx import Document


VLLM_BASE_URL = os.getenv("VLLM_BASE_URL", "http://localhost:11434")
VLLM_MODEL = os.getenv("VLLM_MODEL", "gemma4")

SYSTEM_PROMPT = (
    "あなたはプロの文章校正エディターです。提供されたテキストから"
    "『用語の揺れ』や『明らかな誤字脱字・タイポ』を検出し、"
    "必ずJSON配列のみで返してください。"
    "各要素は id, original, detected_error, suggested_fix, reason を持つこと。"
    "解説文や前置きは返さないでください。"
)


def _post_json(url: str, payload: dict, timeout_seconds: float) -> dict:
    data = json.dumps(payload).encode("utf-8")
    req = request.Request(
        url,
        data=data,
        method="POST",
        headers={"Content-Type": "application/json"},
    )
    with request.urlopen(req, timeout=timeout_seconds) as resp:
        return json.loads(resp.read().decode("utf-8"))


def _get_json(url: str, timeout_seconds: float) -> dict:
    req = request.Request(url, method="GET")
    with request.urlopen(req, timeout=timeout_seconds) as resp:
        text = resp.read().decode("utf-8")
        return json.loads(text) if text.strip() else {}


def check_vllm_availability(timeout_seconds: float = 3.0) -> tuple[bool, str | None, str]:
    """vLLMの疎通確認を行い、利用可否と利用モデルを返す。"""

    base = VLLM_BASE_URL.rstrip("/")
    try:
        models = _get_json(f"{base}/v1/models", timeout_seconds=timeout_seconds)
        data = models.get("data", []) if isinstance(models, dict) else []
        model_id = data[0].get("id") if data else None
        if model_id:
            return True, model_id, "LLM校正機能: 有効"
    except Exception:
        pass

    try:
        _ = request.urlopen(f"{base}/health", timeout=timeout_seconds)
        return True, None, "LLM校正機能: 有効（モデル名未取得）"
    except Exception as exc:
        return (
            False,
            None,
            f"vLLMサーバーへ接続できないため、スタイル統一処理のみを実行しました ({exc.__class__.__name__})",
        )


def _iter_table_paragraphs(container) -> Iterable:
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                yield from _iter_table_paragraphs(cell)


def extract_text_blocks(doc_path: Path) -> list[dict]:
    """本文と表内テキストをID付きで抽出する。"""

    document = Document(str(doc_path))
    blocks: list[dict] = []

    for idx, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            blocks.append({"id": f"p_{idx:03d}", "original": text})

    table_index = 0
    for table in document.tables:
        table_index += 1
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                for para_index, paragraph in enumerate(cell.paragraphs, start=1):
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    blocks.append(
                        {
                            "id": f"t_{table_index:03d}_r_{row_index:03d}_c_{cell_index:03d}_p_{para_index:03d}",
                            "original": text,
                        }
                    )

    # ヘッダー/フッター内の表を取りこぼさないように補完抽出
    for section_index, section in enumerate(document.sections, start=1):
        for para_index, paragraph in enumerate(section.header.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                blocks.append({"id": f"h_{section_index:02d}_{para_index:03d}", "original": text})
        for paragraph in _iter_table_paragraphs(section.header):
            text = paragraph.text.strip()
            if text:
                blocks.append({"id": f"h_{section_index:02d}_t_{len(blocks)+1:03d}", "original": text})
        for para_index, paragraph in enumerate(section.footer.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                blocks.append({"id": f"f_{section_index:02d}_{para_index:03d}", "original": text})
        for paragraph in _iter_table_paragraphs(section.footer):
            text = paragraph.text.strip()
            if text:
                blocks.append({"id": f"f_{section_index:02d}_t_{len(blocks)+1:03d}", "original": text})

    return blocks


def _extract_json_array(text: str) -> list[dict]:
    body = text.strip()
    if body.startswith("```"):
        body = re.sub(r"^```(?:json)?\s*", "", body)
        body = re.sub(r"\s*```$", "", body)

    try:
        parsed = json.loads(body)
        return parsed if isinstance(parsed, list) else []
    except json.JSONDecodeError:
        pass

    match = re.search(r"\[.*\]", body, re.DOTALL)
    if not match:
        return []
    try:
        parsed = json.loads(match.group(0))
        return parsed if isinstance(parsed, list) else []
    except json.JSONDecodeError:
        return []


def request_proofreading(
    blocks: list[dict],
    model_id: str | None,
    timeout_seconds: float = 30.0,
) -> list[dict]:
    """抽出テキストをvLLMへ送り、校正指摘のJSON配列を受け取る。"""

    if not blocks:
        return []

    base = VLLM_BASE_URL.rstrip("/")
    payload = {
        "model": model_id or VLLM_MODEL,
        "temperature": 0,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": json.dumps(blocks, ensure_ascii=False)},
        ],
    }

    response = _post_json(f"{base}/v1/chat/completions", payload, timeout_seconds=timeout_seconds)
    choices = response.get("choices", []) if isinstance(response, dict) else []
    if not choices:
        return []

    message = choices[0].get("message", {})
    content = message.get("content", "")
    return _extract_json_array(content)
