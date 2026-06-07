import copy
import io
import json
import re
import time
from dataclasses import dataclass
from pathlib import Path, PurePosixPath
from typing import Iterable

import jaconv
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from fastapi import UploadFile


MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024
PROCESS_TIMEOUT_SECONDS = 5 * 60
DOWNLOAD_TTL_SECONDS = 30 * 60
DOWNLOAD_CLEANUP_INTERVAL_SECONDS = 5 * 60
DEFAULT_FONTS = [
    "MS Gothic",
    "MS Mincho",
    "Yu Gothic",
    "Meiryo",
    "Arial",
]
UNIFIED_FONT_SIZE_PT = 10.5
SECTION_HEADING_RE = re.compile(
    r"^\s*(?P<prefix>(?:X\d*\.?|\d+(?:\.\d+)*\.?))\s*(?P<title>\S.*)$"
)


@dataclass
class ResultBlob:
    """ダウンロード用ZIPを一時保持するデータ。"""

    filename: str
    data: bytes
    created_at: float


DOWNLOAD_STORE: dict[str, ResultBlob] = {}


def cleanup_expired_downloads(ttl_seconds: int = DOWNLOAD_TTL_SECONDS) -> None:
    """保持期限を過ぎたダウンロードデータを削除する。"""

    now = time.time()
    expired = [k for k, v in DOWNLOAD_STORE.items() if (now - v.created_at) > ttl_seconds]
    for key in expired:
        DOWNLOAD_STORE.pop(key, None)


def sanitize_relative_path(raw_name: str, fallback_name: str) -> Path:
    """アップロードされた相対パスを安全なパスへ正規化する。"""

    candidate = raw_name.strip().replace("\\", "/")
    if not candidate:
        candidate = fallback_name
    posix = PurePosixPath(candidate)
    safe_parts = [p for p in posix.parts if p not in ("", ".", "..")]
    if not safe_parts:
        safe_parts = [Path(fallback_name).name or "upload.docx"]
    return Path(*safe_parts)


def uniquify_path(path: Path, used: set[str]) -> Path:
    """同名ファイルがある場合に連番付きで重複しないパスを返す。"""

    path_str = str(path)
    if path_str not in used:
        used.add(path_str)
        return path

    suffix_index = 1
    stem = path.stem
    suffix = path.suffix
    parent = path.parent
    while True:
        candidate = parent / f"{stem}_{suffix_index}{suffix}"
        candidate_str = str(candidate)
        if candidate_str not in used:
            used.add(candidate_str)
            return candidate
        suffix_index += 1


async def save_upload(upload: UploadFile, destination: Path) -> int:
    """アップロードファイルをサイズ上限を見ながら保存する。"""

    destination.parent.mkdir(parents=True, exist_ok=True)
    total = 0
    with destination.open("wb") as f:
        while True:
            chunk = await upload.read(1024 * 1024)
            if not chunk:
                break
            total += len(chunk)
            if total > MAX_FILE_SIZE_BYTES:
                raise ValueError(f"size limit exceeded: {MAX_FILE_SIZE_BYTES} bytes")
            f.write(chunk)
    await upload.close()
    return total


def is_halfwidth_ascii_alnum_symbol(ch: str) -> bool:
    """文字が半角ASCII英数字記号かどうかを判定する。"""

    code = ord(ch)
    return 0x20 <= code <= 0x7E


def normalize_section_prefix_text(text: str) -> str:
    """見出し接頭辞に含まれる全角英数字記号を半角へ寄せる。"""

    replaced = text.replace("．", ".").replace("Ｘ", "X").replace("ｘ", "x")
    normalized = jaconv.z2h(replaced, kana=False, ascii=True, digit=True)
    return normalized.upper()


def compact_heading_text(text: str) -> str:
    """見出し判定前に改行・全角空白・連続空白を正規化する。"""

    compact = text.replace("\r", " ").replace("\n", " ").replace("\u3000", " ")
    compact = re.sub(r"\s+", " ", compact).strip()
    return normalize_section_prefix_text(compact)


def normalize_heading_text(text: str) -> str:
    """見出し文字列を `接頭辞 + 半角スペース1つ + 本文` に揃える。"""

    compact = compact_heading_text(text)
    match = SECTION_HEADING_RE.match(compact)
    if not match:
        return compact
    prefix = match.group("prefix")
    title = match.group("title")
    return f"{prefix} {title}".strip()


def is_section_heading(paragraph_text: str, heading_allowed: bool) -> bool:
    """段落が見出しとして扱えるかを判定する。"""

    if not heading_allowed:
        return False
    compact = compact_heading_text(paragraph_text)
    return SECTION_HEADING_RE.match(compact) is not None


def _localname(tag: str) -> str:
    """XMLタグから名前空間を除いたローカル名を返す。"""

    return tag.split("}")[-1]


def _paragraph_has_page_break(paragraph: Paragraph) -> bool:
    """段落内に改ページが含まれているかを判定する。"""

    return bool(paragraph._element.xpath('.//w:br[@w:type="page"]'))


def _is_blank_or_pagebreak_paragraph(paragraph: Paragraph) -> bool:
    """文字を持たない空段落または改ページ段落かを判定する。"""

    if paragraph.text.strip():
        return False
    if paragraph._element.xpath(".//w:drawing") or paragraph._element.xpath(".//w:pict"):
        return False
    return True


def normalize_blank_line_before_heading(document: Document) -> None:
    """本文中の見出し直前の区切りを必ず1つだけに揃える。"""

    body = document._body._element
    blocks = list(body.iterchildren())
    idx = 0

    while idx < len(blocks):
        block = blocks[idx]
        if _localname(block.tag) != "p":
            idx += 1
            continue

        paragraph = Paragraph(block, document._body)
        if not is_section_heading(paragraph.text, heading_allowed=True):
            idx += 1
            continue

        back = idx - 1
        separator_indexes: list[int] = []
        # 見出し直前に連続する区切り段落を集める。
        while back >= 0:
            prev_block = blocks[back]
            if _localname(prev_block.tag) != "p":
                break
            prev_paragraph = Paragraph(prev_block, document._body)
            if not _is_blank_or_pagebreak_paragraph(prev_paragraph):
                break
            separator_indexes.append(back)
            back -= 1

        if separator_indexes:
            keep_idx = separator_indexes[0]
            # 改ページ段落がある場合は、それを1つだけ残す。
            for cand_idx in separator_indexes:
                cand_paragraph = Paragraph(blocks[cand_idx], document._body)
                if _paragraph_has_page_break(cand_paragraph):
                    keep_idx = cand_idx
                    break

            for remove_idx in sorted(separator_indexes):
                if remove_idx == keep_idx:
                    continue
                body.remove(blocks[remove_idx])

            blocks = list(body.iterchildren())
            try:
                idx = blocks.index(block)
            except ValueError:
                idx = 0
            idx += 1
            continue

        if idx > 0:
            blank = OxmlElement("w:p")
            block.addprevious(blank)
            blocks = list(body.iterchildren())
            try:
                idx = blocks.index(block)
            except ValueError:
                idx = 0

        idx += 1


def remove_blank_body_line_after_heading(document: Document) -> None:
    """見出し直後に連続する空行を削除する。"""

    body = document._body._element
    blocks = list(body.iterchildren())

    idx = 0
    while idx < len(blocks):
        block = blocks[idx]
        if _localname(block.tag) != "p":
            idx += 1
            continue

        paragraph = Paragraph(block, document._body)
        if not is_section_heading(paragraph.text, heading_allowed=True):
            idx += 1
            continue

        scan = idx + 1
        # 見出し直後の空行だけを消し、本文や表に当たったらそこで止める。
        while scan < len(blocks):
            next_block = blocks[scan]
            if _localname(next_block.tag) != "p":
                break

            next_paragraph = Paragraph(next_block, document._body)
            if next_paragraph.text.strip():
                break

            body.remove(next_block)
            blocks.pop(scan)

        idx += 1


def paragraph_contains_non_text_nodes(paragraph) -> bool:
    """画像などの非テキスト要素を含む段落かを判定する。"""

    for run in paragraph.runs:
        if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
            return True
    return False


def replace_paragraph_text(paragraph, new_text: str) -> None:
    """段落内のRunを入れ替えて段落テキスト全体を書き換える。"""

    p_elem = paragraph._element
    run_elems = list(p_elem.xpath("./w:r"))
    for run_elem in run_elems:
        p_elem.remove(run_elem)
    paragraph.add_run(new_text)


def apply_font(run, font_name: str) -> None:
    """Runにフォント名と統一サイズを設定する。"""

    run.font.name = font_name
    run.font.size = Pt(UNIFIED_FONT_SIZE_PT)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def apply_decoration(run, is_heading: bool) -> None:
    """Runの太字・斜体ルールを適用する。"""

    run.bold = True if is_heading else False
    run.italic = False


def copy_run_properties(src_run, dst_run) -> None:
    """元Runの文字プロパティXMLを新しいRunへ複製する。"""

    src_rpr = src_run._element.rPr
    if src_rpr is None:
        return

    dst_rpr = dst_run._element.get_or_add_rPr()
    for child in list(dst_rpr):
        dst_rpr.remove(child)
    for key, value in src_rpr.attrib.items():
        dst_rpr.set(key, value)
    for child in src_rpr:
        dst_rpr.append(copy.deepcopy(child))


def split_segments(text: str, force_halfwidth_ascii: bool = False) -> list[tuple[str, bool]]:
    """文字種ごとに分割し、必要なら英数字記号を半角化する。"""

    if not text:
        return []

    converted = jaconv.h2z(text, kana=True, ascii=False, digit=False)
    if force_halfwidth_ascii:
        converted = jaconv.z2h(converted, kana=False, ascii=True, digit=True)
    segments: list[tuple[str, bool]] = []
    buff = [converted[0]]
    current_ascii = is_halfwidth_ascii_alnum_symbol(converted[0])
    for ch in converted[1:]:
        is_ascii = is_halfwidth_ascii_alnum_symbol(ch)
        if is_ascii == current_ascii:
            buff.append(ch)
            continue
        segments.append(("".join(buff), current_ascii))
        buff = [ch]
        current_ascii = is_ascii
    segments.append(("".join(buff), current_ascii))
    return segments


def normalize_paragraph(
    paragraph,
    font_a: str,
    font_b: str,
    heading_allowed: bool,
    force_halfwidth_ascii: bool = False,
) -> None:
    """段落単位で見出し正規化、文字変換、装飾統一を行う。"""

    heading = is_section_heading(paragraph.text, heading_allowed=heading_allowed)
    if heading and not paragraph_contains_non_text_nodes(paragraph):
        # 見出し文字列は一度だけ正規化して、接頭辞と空白を整える。
        normalized = normalize_heading_text(paragraph.text)
        if normalized and normalized != paragraph.text:
            replace_paragraph_text(paragraph, normalized)

    runs = list(paragraph.runs)
    for run in runs:
        original_text = run.text
        segments = split_segments(original_text, force_halfwidth_ascii=force_halfwidth_ascii)
        if not segments:
            apply_decoration(run, is_heading=heading)
            run.font.size = Pt(UNIFIED_FONT_SIZE_PT)
            continue

        first_text, first_is_ascii = segments[0]
        run.text = first_text
        apply_font(run, font_a if first_is_ascii else font_b)
        apply_decoration(run, is_heading=heading)

        previous = run
        for text, is_ascii in segments[1:]:
            new_run = paragraph.add_run(text)
            copy_run_properties(run, new_run)
            apply_font(new_run, font_a if is_ascii else font_b)
            apply_decoration(new_run, is_heading=heading)
            previous._element.addnext(new_run._element)
            previous = new_run


def iter_table_paragraphs_from_container(container) -> Iterable:
    """コンテナ配下の表セル内段落を再帰的に列挙する。"""

    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                yield from iter_table_paragraphs_from_container(cell)


def normalize_document(input_path: Path, output_path: Path, font_a: str, font_b: str) -> None:
    """Word文書全体に統一ルールを適用して保存する。"""

    document = Document(str(input_path))

    normalize_blank_line_before_heading(document)
    remove_blank_body_line_after_heading(document)

    for paragraph in document.paragraphs:
        normalize_paragraph(paragraph, font_a, font_b, heading_allowed=True)
    for paragraph in iter_table_paragraphs_from_container(document):
        normalize_paragraph(
            paragraph,
            font_a,
            font_b,
            heading_allowed=False,
            force_halfwidth_ascii=True,
        )

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            normalize_paragraph(paragraph, font_a, font_b, heading_allowed=False)
        for paragraph in iter_table_paragraphs_from_container(section.header):
            normalize_paragraph(
                paragraph,
                font_a,
                font_b,
                heading_allowed=False,
                force_halfwidth_ascii=True,
            )
        for paragraph in section.footer.paragraphs:
            normalize_paragraph(paragraph, font_a, font_b, heading_allowed=False)
        for paragraph in iter_table_paragraphs_from_container(section.footer):
            normalize_paragraph(
                paragraph,
                font_a,
                font_b,
                heading_allowed=False,
                force_halfwidth_ascii=True,
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def build_zip_bytes(
    output_root: Path,
    failed_entries: list[tuple[str, str]],
    llm_report: list[dict] | None = None,
) -> bytes:
    """変換済みファイル群と失敗一覧をZIPバイト列にまとめる。"""

    buffer = io.BytesIO()

    import zipfile

    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file_path in sorted(output_root.rglob("*")):
            if file_path.is_file():
                zf.write(file_path, arcname=str(file_path.relative_to(output_root)))

        if failed_entries:
            lines = ["path\treason"]
            lines.extend(f"{path}\t{reason}" for path, reason in failed_entries)
            zf.writestr("faield.txt", "\n".join(lines) + "\n")

        if llm_report:
            zf.writestr("llm_findings.json", json.dumps(llm_report, ensure_ascii=False, indent=2))

    buffer.seek(0)
    return buffer.getvalue()
